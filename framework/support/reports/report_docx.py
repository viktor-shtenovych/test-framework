import os
from typing import List

from docx import Document
from docx.shared import Inches, Cm

from framework.support.reports.report_data import TestCase


def generate_report(test_cases: List[TestCase], path: str) -> None:
    """
    Generate a test report in a Word document.

    Args:
        test_cases (List[TestCase]): The test cases to include in the report.
        path (str): The path to save the report to.
    """
    document = Document()

    document.add_heading("Test report", 0)

    document.add_paragraph("Footswitch test execution report")

    document.add_page_break()  # type: ignore
    document.add_heading("Test Procedures and results", 1)

    group = None

    for test in test_cases:
        if test.skipped:
            continue

        if group != test.group:
            group = test.group
            document.add_heading(group, 2)

        summary = document.add_paragraph()
        summary.add_run(test.test_id).bold = True
        summary.add_run(test.summary).italic = True

        table = document.add_table(rows=1, cols=3)
        table.allow_autofit = True
        table.columns[0].width = Cm(1.3)
        table.columns[1].width = Cm(7)
        table.columns[2].width = Cm(7)

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Step"
        hdr_cells[1].text = "Action"
        hdr_cells[2].text = "Expected result"

        for number, step in enumerate(test.steps):
            row_cells = table.add_row().cells
            row_cells[0].text = str(number)
            row_cells[1].text = step.action
            row_cells[2].text = ""
            for verification in step.verifications:
                p = row_cells[2].add_paragraph("")
                r = p.add_run()
                if verification[0]:
                    r.add_picture(
                        rf"{os.path.dirname(__file__)}/resources/pass.png",
                        width=Inches(0.1),
                    )
                else:
                    r.add_picture(
                        rf"{os.path.dirname(__file__)}/resources/fail.png",
                        width=Inches(0.1),
                    )
                p.add_run(verification[1])
        document.add_paragraph().add_run().add_break()

    document.save(path)
